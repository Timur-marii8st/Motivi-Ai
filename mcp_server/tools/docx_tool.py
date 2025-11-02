from __future__ import annotations
from pathlib import Path
from datetime import datetime, timezone
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger
import asyncio

from ..config import mcp_settings

async def create_docx(
    chat_id: int,
    title: str,
    sections: list,
    footer: str | None = None,
) -> Path:
    """
    Generate a .docx file with title, sections, and optional footer.
    """
    doc = Document()
    
    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sections
    for sec in sections:
        doc.add_heading(sec["heading"], level=1)
        doc.add_paragraph(sec["content"])
    
    # Footer
    if footer:
        doc.add_paragraph() # Adds a space before the footer
        footer_para = doc.add_paragraph(footer)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.runs[0]
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    filename = f"plan_{chat_id}_{timestamp}.docx"
    file_path = Path(mcp_settings.TEMP_FILES_DIR) / filename
    
    # Run the blocking doc.save() operation in a separate thread
    loop = asyncio.get_running_loop()
    await loop.run_in_executor(None, doc.save, str(file_path))
    
    logger.info("Created docx: {}", file_path)
    return file_path